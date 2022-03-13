from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(16)
style.font.color.rgb = RGBColor(255, 0, 0)

title = doc.add_heading('my title', 5)
title.add_run('123') # 追加内容

title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
title.style.font.size = Pt(20)
title.italic = True
title.bold = True

p = doc.add_paragraph('欢迎学习python!')
p.add_run('\n这是word 段落增加')

image = doc.add_picture('logo2020.png', width=Inches(2), height=Inches(2))

# 换页
doc.add_page_break()

table_title = ["name", "age", "sex"]
table = doc.add_table(rows=1, cols=3)
row_cells = table.rows[0].cells
row_cells[0].text = table_title[0]
row_cells[1].text = table_title[1]
row_cells[2].text = table_title[2]

data = [
    ('xiaomu', '10', 'man'),
    ('dewei', '34', 'man'),
    ('xiaoman', '18', 'women')
]
for d in data:
    row_cells = table.add_row().cells
    row_cells[0].text = d[0]
    row_cells[1].text = d[1]
    row_cells[2].text = d[2]

doc.save('mytest.docx')