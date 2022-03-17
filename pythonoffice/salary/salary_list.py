# coding:utf-8

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
import xlrd
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.header import Header
import glob
import smtplib

# 读取excel表格的内容，将excel的数据用列表返回
def excel_read(path):
    book = xlrd.open_workbook(path)
    sheet1 = book.sheet_by_name('研发部工资表')
    content = []
    for row in sheet1.get_rows():
        row_data = []
        for cell in row:
            row_data.append(cell.value)
        content.append(row_data)
    return content


excel_data = excel_read('salary.xlsx')
doc = Document()

# 设置工资单word的样式
def word_style():
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')  # 中文字体

    title = doc.add_heading("", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _t = title.add_run('5月工资单')
    _t.font.name = '微软雅黑'
    _t.element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    _t.bold = True


# 将表格的内容读取出来，并且按照设置好的样式，生成每个人的工资单word，并保存
def docx_write():
    word_style()
    table = doc.add_table(rows=2, cols=10, style='Light List Accent 5')
    table_header_list = ['姓名', '性别', '部门', '基本工资', '绩效工资', '餐补', '社保', '公积金', '代扣个税', '实发工资']
    table_header = table.rows[0].cells
    for i in range(len(table_header_list)):
        table_header[i].text = table_header_list[i]

    for data in excel_data[1:]:
        row_cells = table.rows[1].cells
        for index, value in enumerate(data[0:-1]):
            row_cells[index].text = str(value)
        doc.save('{}.docx'.format(data[0]))


# print(excel_read('salary.xlsx'))
# docx_write()

# 发邮件，内容包含html内容，和word附件
def send_email(mail_host, mail_user, mail_pass):
    # sender 发送邮箱我这边给空值啊
    sender = ''
    receiver_emails = []
    # 把所有的邮箱放入列表
    for data in excel_data[1:]:
        receiver_emails.append(data[len(data)-1])

    for receiver_names in excel_data[1:]:
        message = MIMEMultipart()
        message['From'] = Header(sender)
        message['Subject'] = Header('工资单', 'utf-8')
        path = glob.os.path.join(glob.os.getcwd(), '{}.docx'.format(receiver_names[0]))
        attr = MIMEText(open(path, 'rb').read(), 'base64', 'utf-8')
        attr['Content-Type'] = 'application/octet-stream'
        # 附件名称为中文时的写法
        attr.add_header("Content-Disposition", "attachment", filename=("gbk", "", receiver_names[0] + ".docx"))
        # 附件名称非中文时的写法
        # attr["Content-Disposition"] = 'attachment; filename="salary.docx")'
        message.attach(attr)
        msg_content = """
        <p style="font-weight:bold;">亲爱的
        """ + receiver_names[0] + """
        ：</p>
        <p style="text-indent:2em;">感谢您为公司做出的贡献！</p>
        <p style="text-indent:2em;">本月工资已经到账，请注意查收。如未收到，请您在5个工作日内联系人力资源部。</p>
        <p style="text-indent:2em;">您的工资单详情请查看邮件附件！</p>
        <p style="text-indent:2em;">如有异议请联系人力资源部。</p></p>
        <br /><br /><br />
        <p style="text-align:right;">人力资源部<br />2021年5月31日</p>
        """

        message.attach(MIMEText(msg_content, 'html', 'utf-8'))
        try:
            smtpobj = smtplib.SMTP()
            smtpobj.connect(mail_host, 25)  # 端口号
            smtpobj.login(mail_user, mail_pass)
            smtpobj.sendmail(sender, receiver_names[10], message.as_string())
        except smtplib.SMTPException as e:
            print(e)

if __name__ == '__main__':
    docx_write()
    mail_host = 'smtp.qq.com'
# 账号密码涉及隐私，这边我给空值，老师测试的时候可以用自己的账号，上面还有sender也隐藏掉了
    mail_user = ''
    mail_pass = ''
    send_email(mail_host, mail_user, mail_pass)




