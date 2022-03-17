from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
doc = Document()

title = doc.add_heading('慕课网是谁', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

p = doc.add_paragraph('我们是IT教育行业的造梦者，也是前沿技术内容的创造者和传播者！')
p.add_run('\n体系课：')
# p.add_run('\n1.python全栈工程师')
# p.add_run('\n2.java工程师')
# p.add_run('\n3.前端工程师')
doc.add_paragraph('python全栈工程师', style='List Number')
doc.add_paragraph('Java工程师', style='List Number')
doc.add_paragraph('前端工程师', style='List Number')

p.style.font.name = '微软雅黑'
p.style.font.size = Pt(15)

doc.add_page_break()

title = doc.add_heading('学生信息', 0)
table_title = ['序号', '姓名', '年龄', '身高']
table = doc.add_table(rows=1, cols=4)
one_row_cells = table.rows[0].cells
one_row_cells[0].text = table_title[0]
one_row_cells[1].text = table_title[1]
one_row_cells[2].text = table_title[2]
one_row_cells[3].text = table_title[3]

data = [['1', '张三', "20", "174"], ["2", "李四", '19', '167']]
for d in data:
    row_add = table.add_row().cells
    row_add[0].text = d[0]
    row_add[1].text = d[1]
    row_add[2].text = d[2]
    row_add[3].text = d[3]

doc.save('demo.docx')