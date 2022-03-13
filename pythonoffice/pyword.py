from docx import Document


doc = Document('文本.docx')

for p in doc.paragraphs:
    print(p.text)

for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            print(cell.text)