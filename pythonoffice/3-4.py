from docx import Document

doc = Document('文本文档.docx')

print('段落数：{}'.format(len(doc.paragraphs)))
p_list = []
for p in doc.paragraphs:
    p_list.append(p)

for index, p in enumerate(p_list, start=1):
    print('第{}段的内容是：{}'.format(index, p.text))